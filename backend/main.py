from fastapi import FastAPI, HTTPException, Depends, UploadFile, File
from fastapi.middleware.cors import CORSMiddleware
from sqlalchemy import create_engine, Column, Integer, String, Text
from sqlalchemy.ext.declarative import declarative_base
from sqlalchemy.orm import sessionmaker, Session
from datetime import datetime
from pathlib import Path
import chromadb
import fitz  # PyMuPDF for PDF processing
import docx  # python-docx for DOCX processing
import requests
import json
import asyncio
from typing import List

# Database setup
SQLALCHEMY_DATABASE_URL = "sqlite:///./test.db"
engine = create_engine(SQLALCHEMY_DATABASE_URL, connect_args={"check_same_thread": False})
SessionLocal = sessionmaker(autocommit=False, autoflush=False, bind=engine)
Base = declarative_base()

# Models
class Note(Base):
    __tablename__ = "notes"
    id = Column(Integer, primary_key=True, index=True)
    title = Column(String)
    content = Column(Text)
    folder_id = Column(String)
    created_at = Column(String, default=lambda: datetime.now().isoformat())
    updated_at = Column(String, default=lambda: datetime.now().isoformat())

class Document(Base):
    __tablename__ = "documents"
    id = Column(Integer, primary_key=True, index=True)
    filename = Column(String)
    content = Column(Text)
    created_at = Column(String, default=lambda: datetime.now().isoformat())

# Create tables
Base.metadata.create_all(bind=engine)

# Vector store setup
chroma_client = chromadb.Client()
collection = chroma_client.create_collection(name="documents")

app = FastAPI()

# CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:8080"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Dependency to get DB session
def get_db():
    db = SessionLocal()
    try:
        yield db
    finally:
        db.close()

async def process_document(file_path: Path) -> str:
    """Process document and extract text asynchronously."""
    if file_path.suffix == ".pdf":
        doc = fitz.open(file_path)
        text = ""
        for page in doc:
            text += page.get_text()
        return text
    elif file_path.suffix in [".doc", ".docx"]:
        doc = docx.Document(file_path)
        return "\n".join(para.text for para in doc.paragraphs)
    elif file_path.suffix == ".txt":
        return file_path.read_text()
    else:
        raise ValueError(f"Unsupported file type: {file_path.suffix}")

@app.post("/upload")
async def upload_files(files: List[UploadFile] = File(...), db: Session = Depends(get_db)):
    results = []
    upload_dir = Path("uploads")
    upload_dir.mkdir(exist_ok=True)

    for file in files:
        file_path = upload_dir / file.filename
        
        # Save file
        content = await file.read()
        with file_path.open("wb") as f:
            f.write(content)

        # Process document
        try:
            text = await process_document(file_path)
            
            # Save to database
            doc = Document(filename=file.filename, content=text)
            db.add(doc)
            
            # Add to vector store
            collection.add(
                documents=[text],
                metadatas=[{"source": file.filename}],
                ids=[f"doc_{datetime.now().timestamp()}"]
            )
            
            results.append({"filename": file.filename, "status": "success"})
        except Exception as e:
            results.append({"filename": file.filename, "status": "error", "error": str(e)})

    db.commit()
    return results

@app.post("/notes")
async def create_note(title: str, content: str, folder_id: str, db: Session = Depends(get_db)):
    note = Note(title=title, content=content, folder_id=folder_id)
    db.add(note)
    db.commit()
    db.refresh(note)
    return note

@app.get("/notes")
async def get_notes(db: Session = Depends(get_db)):
    return db.query(Note).all()

@app.post("/chat")
async def chat(message: str, db: Session = Depends(get_db)):
    # Search relevant documents
    results = collection.query(
        query_texts=[message],
        n_results=3
    )

    context = "\n".join(results["documents"][0]) if results.get("documents") else ""
    prompt = f"Context: {context}\n\nQuestion: {message}\n\nAnswer:"

    # Call Ollama API
    try:
        ollama_url = "http://localhost:11434/api/generate"
        payload = {
            "model": "llama2",
            "prompt": prompt,
            "stream": False
        }
        response = await asyncio.get_running_loop().run_in_executor(
            None,
            lambda: requests.post(ollama_url, json=payload)
        )
        
        if response.status_code != 200:
            raise HTTPException(status_code=500, detail="Failed to get response from LLM")
            
        response_json = response.json()
        return {"response": response_json.get("response", "No response generated.")}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8090)
